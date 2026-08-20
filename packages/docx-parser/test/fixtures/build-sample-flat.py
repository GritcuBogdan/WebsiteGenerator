"""Regenerates sample-flat.docx, the fixture parse-docx-flat.test.ts runs
against — a docx where every subsection heading, including "Navigation" and
"FAQ", was left as a plain paragraph instead of Heading 3 (a real defect
seen in a translated locale copy: the whole document had zero Heading 3
paragraphs). Exercises promote_missing_heading3() in casinoParser.py.

Run from packages/docx-parser: .venv/Scripts/python test/fixtures/build-sample-flat.py
"""

from pathlib import Path

from docx import Document

OUTPUT_PATH = Path(__file__).parent / "sample-flat.docx"


def build():
    document = Document()

    # --- Page 1: home --- (Heading 2 still present - only Heading 3 is missing)
    document.add_heading("SEITE 1 - Home", level=2)
    document.add_paragraph("Meta Title: Sample Casino - Home")
    document.add_paragraph("H1: Sample Casino")
    document.add_paragraph("Welcome to Sample Casino, the best place to play.")
    document.add_paragraph("Hero Banner: Get up to $500 bonus today!")

    document.add_paragraph("Navigation")
    document.add_paragraph("About Us")
    document.add_paragraph("Questions People Ask")

    # "H2:" text prefix instead of a real Heading 2/3 style - seen in a
    # second real defective docx alongside the missing-Heading-3 issue.
    document.add_paragraph("H2: About Us")
    document.add_paragraph("We are a great casino with many games.")

    # Untranslatable heading ("Questions People Ask" isn't in any hardcoded
    # FAQ-keyword list) - must be detected structurally, from the Q/A shape
    # of its own paragraphs, not from its title.
    document.add_paragraph("H2: Questions People Ask")
    # Question and answer merged onto one paragraph, as seen in the real
    # defective docx, instead of two separate paragraphs.
    document.add_paragraph("Is this safe? Yes, absolutely safe and secure.")
    document.add_paragraph("How do I withdraw? Go to the withdrawal page and follow the steps.")

    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
