"""Regenerates sample.docx, the fixture parse-docx.test.ts runs against.

Not part of the pipeline itself — a one-off content generator, kept so the
fixture can be regenerated/extended without hand-editing a .docx in Word.
Run from packages/docx-parser: .venv/Scripts/python test/fixtures/build-sample.py
"""

from pathlib import Path

from docx import Document

OUTPUT_PATH = Path(__file__).parent / "sample.docx"


def build():
    document = Document()

    # --- Page 1: home ---
    document.add_heading("SEITE 1 - Home", level=2)
    document.add_paragraph("Meta Title: Sample Casino - Home")
    document.add_paragraph("Meta Description: A sample casino for testing.")
    document.add_paragraph("H1: Sample Casino")
    document.add_paragraph("Welcome to Sample Casino, the best place to play.")
    # A second, unrelated paragraph left over from copy that didn't get its
    # own Heading 3 - regression coverage for the "overflow" split below.
    document.add_paragraph("Extra trust copy that is too long to fit in the hero.")
    document.add_paragraph("Hero Banner: Get up to $500 bonus today!")

    # Styled as a plain paragraph, not Heading 3 - regression coverage for
    # style-agnostic navigation-marker detection.
    document.add_paragraph("Navigation")
    document.add_paragraph("Bonus")
    document.add_paragraph("Slots")

    document.add_heading("About Us", level=3)
    document.add_paragraph("We are a great casino with many games.")

    document.add_heading("FAQ", level=3)
    document.add_paragraph("Is this safe?")
    document.add_paragraph("Yes, absolutely safe and secure.")
    document.add_paragraph("How do I withdraw?")
    document.add_paragraph("Go to the withdrawal page and follow the steps.")

    # --- Page 2: bonus ---
    document.add_heading("SEITE 2 - Bonus", level=2)
    document.add_paragraph("Meta Title: Sample Casino - Bonus")
    document.add_paragraph("H1: Bonus Page")

    document.add_heading("Bonus Details", level=3)
    document.add_paragraph("Here are the details:")
    document.add_paragraph("First deposit bonus")
    document.add_paragraph("Second deposit bonus")

    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Bonus Type"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "Welcome"
    table.cell(1, 1).text = "$500"
    table.cell(2, 0).text = "Reload"
    table.cell(2, 1).text = "$100"

    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
