"""Builds two throwaway docx fixtures (en + de) covering the full mandatory
casino-v1-rules page set (index/login/application/withdrawal/review/bonus
family/slots/legal pages), used to exercise casino-v2's build against
sites/casino-v2-fixture. Not part of the pipeline itself - a one-off content
generator, same pattern as build-sample.py.

Run from packages/docx-parser: .venv/Scripts/python test/fixtures/build-casino-v2-fixture.py
"""

from pathlib import Path

from docx import Document

OUTPUT_DIR = Path(__file__).parent


def add_page(document, marker, meta_title, meta_description, h1, intro_paragraphs, banner_text=None, nav=None):
    document.add_heading(marker, level=2)
    document.add_paragraph(f"Meta Title: {meta_title}")
    document.add_paragraph(f"Meta Description: {meta_description}")
    document.add_paragraph(f"H1: {h1}")
    for paragraph in intro_paragraphs:
        document.add_paragraph(paragraph)
    if banner_text:
        document.add_paragraph(f"Hero Banner: {banner_text}")
    if nav:
        document.add_heading("Navigation", level=3)
        for item in nav:
            document.add_paragraph(item)


def add_section(document, title, paragraphs):
    document.add_heading(title, level=3)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)


def add_faq(document, title, pairs):
    document.add_heading(title, level=3)
    for question, answer in pairs:
        document.add_paragraph(question)
        document.add_paragraph(answer)


def build_en():
    document = Document()

    add_page(
        document,
        "PAGE 1 - Home",
        "Sample Casino V2 - Home",
        "A long-running online casino review with bonuses, slots, and payment guides.",
        "Sample Casino V2",
        [
            "Welcome to Sample Casino V2, an independent review covering everything a new "
            "player needs before signing up: bonuses, banking, and our full slots library. "
            "This introduction is intentionally long to exercise hero text wrapping across "
            "narrow mobile widths, tablet widths, and wide desktop layouts without breaking "
            "the layout or clipping any content.",
        ],
        banner_text="Get up to $1000 bonus + 200 free spins today!",
        nav=["Bonus", "Slots"],
    )
    add_section(
        document,
        "About Us",
        ["We independently review online casinos so players can compare bonuses, payout speed, and game libraries before signing up."],
    )
    add_faq(
        document,
        "FAQ",
        [
            ("Is this casino safe?", "Yes, it holds a valid gaming license and uses industry-standard encryption."),
            ("How long do withdrawals take?", "Most withdrawals are processed within 24 to 48 hours, depending on the payment method."),
        ],
    )

    add_page(document, "PAGE 2 - Login", "Sample Casino V2 - Login", "Log in to your account.", "Login", ["Access your account to manage deposits, bonuses, and game history."])
    add_section(document, "Login Help", ["If you forgot your password, use the reset link on the login form."])

    add_page(document, "PAGE 3 - Registration", "Sample Casino V2 - Registration", "Create a new account.", "Registration", ["Sign up in under two minutes with just an email and password."])
    add_section(document, "Why Sign Up", ["New players get a welcome bonus on their first deposit."])

    add_page(document, "PAGE 4 - Withdrawal", "Sample Casino V2 - Withdrawal", "How to withdraw your winnings.", "Withdrawal", ["Withdrawals are processed daily and paid out via your original deposit method wherever possible."])
    add_section(
        document,
        "Withdrawal Steps",
        [
            "Follow these steps to request a withdrawal:",
            "Log in to your account",
            "Open the cashier and select withdraw",
            "Choose your payment method",
            "Confirm the amount and submit",
        ],
    )

    add_page(document, "PAGE 5 - Review", "Sample Casino V2 - Review", "Our full review.", "Review", ["Our editorial team tested deposits, withdrawals, game fairness, and customer support."])
    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Category"
    table.cell(0, 1).text = "Rating"
    table.cell(1, 0).text = "Game Selection"
    table.cell(1, 1).text = "9.2 / 10"
    table.cell(2, 0).text = "Payout Speed"
    table.cell(2, 1).text = "8.7 / 10"
    table.cell(3, 0).text = "Customer Support"
    table.cell(3, 1).text = "9.0 / 10"

    add_page(document, "PAGE 6 - Bonus", "Sample Casino V2 - Bonus", "All current bonus offers.", "Bonus Offers", ["Browse every active promotion, from welcome bonuses to reload offers."])

    add_page(document, "PAGE 7 - No Deposit Bonus", "Sample Casino V2 - No Deposit Bonus", "Bonus with no deposit required.", "No Deposit Bonus", ["Claim a small bonus just for signing up, no deposit needed."])

    add_page(document, "PAGE 8 - Free Spins", "Sample Casino V2 - Free Spins", "Free spins offers.", "Free Spins", ["Get free spins on selected slot titles every week."])

    add_page(document, "PAGE 9 - Promo Code", "Sample Casino V2 - Promo Code", "Current promo codes.", "Promo Codes", ["Use these codes at checkout to unlock extra bonus value."])

    add_page(document, "PAGE 10 - Slots", "Sample Casino V2 - Slots", "Our favorite slots.", "Slots", ["A hand-picked library of the best-performing slot titles this month."])

    add_page(document, "PAGE 11 - Responsible Gaming", "Sample Casino V2 - Responsible Gaming", "Play responsibly.", "Responsible Gaming", ["Set deposit limits and self-exclusion tools are available in your account settings."])

    add_page(document, "PAGE 12 - Privacy Policy", "Sample Casino V2 - Privacy Policy", "How we handle your data.", "Privacy Policy", ["We only collect the data required to operate your account and never sell it to third parties."])

    add_page(document, "PAGE 13 - Cookies Policy", "Sample Casino V2 - Cookies Policy", "Our cookie usage.", "Cookies Policy", ["This site uses cookies for essential functionality and anonymized analytics."])

    add_page(document, "PAGE 14 - Betting Rules", "Sample Casino V2 - Betting Rules", "Game rules.", "Betting Rules", ["Standard game rules apply to every title unless otherwise stated on the game page."])

    add_page(document, "PAGE 15 - Terms and Conditions", "Sample Casino V2 - Terms & Conditions", "Full terms.", "Terms & Conditions", ["By using this site you agree to the terms outlined here in full."])

    add_page(document, "PAGE 16 - Contact", "Sample Casino V2 - Contact", "Get in touch.", "Contacts", ["Reach our support team by email or live chat, available around the clock."])

    document.save(OUTPUT_DIR / "casino-v2-fixture.en.docx")
    print(f"Wrote {OUTPUT_DIR / 'casino-v2-fixture.en.docx'}")


def build_de():
    document = Document()

    add_page(
        document,
        "SEITE 1 - Startseite",
        "Beispiel Casino V2 - Startseite",
        "Eine unabhängige Bewertung mit Boni, Slots und Zahlungsanleitungen.",
        "Beispiel Casino V2",
        ["Willkommen bei Beispiel Casino V2, einer unabhängigen Bewertung für neue Spieler."],
        banner_text="Erhalte bis zu 1000€ Bonus + 200 Freispiele!",
        nav=["Bonus", "Slots"],
    )
    add_faq(
        document,
        "Häufig gestellte Fragen",
        [
            ("Ist dieses Casino sicher?", "Ja, es verfügt über eine gültige Glücksspiellizenz."),
            ("Wie lange dauern Auszahlungen?", "Die meisten Auszahlungen werden innerhalb von 24 bis 48 Stunden bearbeitet."),
        ],
    )

    add_page(document, "SEITE 2 - Anmeldung", "Beispiel Casino V2 - Anmeldung", "Melde dich an.", "Anmeldung", ["Greife auf dein Konto zu, um Einzahlungen und Boni zu verwalten."])

    add_page(document, "SEITE 3 - Registrierung", "Beispiel Casino V2 - Registrierung", "Erstelle ein Konto.", "Registrierung", ["Registriere dich in weniger als zwei Minuten."])

    add_page(document, "SEITE 4 - Auszahlung", "Beispiel Casino V2 - Auszahlung", "So zahlst du aus.", "Auszahlung", ["Auszahlungen werden täglich bearbeitet."])

    add_page(document, "SEITE 5 - Bewertung", "Beispiel Casino V2 - Bewertung", "Unsere Bewertung.", "Bewertung", ["Unser Team hat Einzahlungen, Auszahlungen und den Support getestet."])

    add_page(document, "SEITE 6 - Bonus", "Beispiel Casino V2 - Bonus", "Alle aktuellen Bonusangebote.", "Bonusangebote", ["Durchstöbere alle aktiven Aktionen."])

    add_page(document, "SEITE 7 - Bonus ohne Einzahlung", "Beispiel Casino V2 - Bonus ohne Einzahlung", "Bonus ohne Einzahlung.", "Bonus ohne Einzahlung", ["Sichere dir einen kleinen Bonus ganz ohne Einzahlung."])

    add_page(document, "SEITE 8 - Freispiele", "Beispiel Casino V2 - Freispiele", "Freispiel-Angebote.", "Freispiele", ["Erhalte jede Woche Freispiele auf ausgewählte Slots."])

    add_page(document, "SEITE 9 - Aktionscode", "Beispiel Casino V2 - Aktionscode", "Aktuelle Aktionscodes.", "Aktionscodes", ["Nutze diese Codes an der Kasse."])

    add_page(document, "SEITE 10 - Spiele", "Beispiel Casino V2 - Slots", "Unsere Lieblingsslots.", "Slots", ["Eine handverlesene Auswahl der besten Slot-Titel."])

    add_page(document, "SEITE 11 - Verantwortungsvolles Spielen", "Beispiel Casino V2 - Verantwortungsvolles Spielen", "Spiele verantwortungsvoll.", "Verantwortungsvolles Spielen", ["Einzahlungslimits und Selbstausschluss sind in den Kontoeinstellungen verfügbar."])

    add_page(document, "SEITE 12 - Datenschutz", "Beispiel Casino V2 - Datenschutz", "Wie wir deine Daten behandeln.", "Datenschutzerklärung", ["Wir sammeln nur die Daten, die zum Betrieb deines Kontos erforderlich sind."])

    add_page(document, "SEITE 13 - Cookie-Richtlinie", "Beispiel Casino V2 - Cookie-Richtlinie", "Unsere Cookie-Nutzung.", "Cookie-Richtlinie", ["Diese Website verwendet Cookies für wesentliche Funktionen."])

    add_page(document, "SEITE 14 - Spielregeln", "Beispiel Casino V2 - Spielregeln", "Spielregeln.", "Spielregeln", ["Für jeden Titel gelten die Standardspielregeln."])

    add_page(document, "SEITE 15 - Allgemeine Geschäftsbedingungen", "Beispiel Casino V2 - AGB", "Vollständige AGB.", "Allgemeine Geschäftsbedingungen", ["Mit der Nutzung dieser Seite stimmst du den hier genannten Bedingungen zu."])

    add_page(document, "SEITE 16 - Kontakt", "Beispiel Casino V2 - Kontakt", "Kontaktiere uns.", "Kontakt", ["Erreiche unser Support-Team jederzeit per E-Mail oder Live-Chat."])

    document.save(OUTPUT_DIR / "casino-v2-fixture.de.docx")
    print(f"Wrote {OUTPUT_DIR / 'casino-v2-fixture.de.docx'}")


if __name__ == "__main__":
    build_en()
    build_de()
